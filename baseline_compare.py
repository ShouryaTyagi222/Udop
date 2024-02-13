# python udop_baseline_compare.py -i '/data/circulars/DATA/split_circulars/SplitCircularsv2/first_page' -o '/data/circulars/DATA/Models/CircularsV1/docvqa.json'

# Imports
import json
import os
from transformers import AutoProcessor, UdopForConditionalGeneration
from PIL import Image
from docx import Document
from docx.shared import Inches
import argparse
import warnings
warnings.filterwarnings("ignore")

q_map={
    'q1':"Which organization issued this given circular?",
    'q2': "What is the Address of the Issuing Authority of the given Circular?",
    'q3': "What is the Serial No./ID of the Given Circular?",
    'q4': "What is the Date of Issuance of the Circular?",
    'q5': "What is the Subject of the given Circular?",
    'q6': "Who has this circular been addressed to?",
    'q7': "To Whom has the circular been forwarded to?",
    'q8': "Who Has Forwarded This Circular?",
    'q9': "What is the Designation of the Person who Forwarded this Circular?",
    'q10': "Who has signed the Given Circular?",
    'q11': "What is the Designation of the Person who Signed this Circular?"
}



def infer(args):
    with open(args.data_file,'r') as f:
        data=json.load(f)
    
    processor = AutoProcessor.from_pretrained("nielsr/udop-large", apply_ocr=True)
    model = UdopForConditionalGeneration.from_pretrained("nielsr/udop-large")
    
    doc=Document()
    i=0
    
    for da in data:
        img_name=da['file_name']
        doc.add_heading(img_name, 1)
        img_path=os.path.join(args.image_dir,img_name)
        doc.add_picture(img_path,width=Inches(3)).alignment=1
        image = Image.open(img_path).convert("RGB")
        annotations=da['annotations']
        for annotation in annotations:
            try:
                if annotation['type']=='textarea':
                    value=annotation['value']['text'][0]
                    q=q_map[annotation['to_name']]
                    prompt = f"Question answering. {q}"
                    encoding = processor(images=image, text=prompt, return_tensors="pt")
                    predicted_ids = model.generate(**encoding)
                    output=processor.batch_decode(predicted_ids, skip_special_tokens=True)[0]
                    print(i)
                    print('Question :',q)
                    print(f'Orignial Answer : {value}')
                    print('Udop :',output)
                    p=doc.add_paragraph(f'Question : {q}\n >>>Orignial Answer : {value} \n >>>Udop : {output}\n')
                    i+=1
                if i==10:
                    doc.save('infer_results.docx')
                    return 0
            except Exception as e:
                print(e)
        print()
    doc.save('infer_results.docx')

def parse_args():
    parser = argparse.ArgumentParser(description="Model INFER", formatter_class=argparse.ArgumentDefaultsHelpFormatter)

    parser.add_argument("-i", "--image_dir", type=str, default=None, help="path to the image dir")
    parser.add_argument("-o", "--data_file", type=str, default="OUTPUT", help="path to the processed annotation json data file")
    args = parser.parse_args()
    return args


if __name__ == "__main__":
    args = parse_args()
    infer(args)